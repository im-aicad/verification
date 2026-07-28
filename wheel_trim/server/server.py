"""Wheelhouse regulation verification web server."""

from __future__ import annotations

import asyncio
import json
import logging
import os
import re
import socket
import shutil
import subprocess
import sys
import tempfile
import traceback
import zipfile
from contextlib import asynccontextmanager
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape

from fastapi import FastAPI, File, Form, HTTPException, UploadFile, WebSocket, WebSocketDisconnect
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles


IS_FROZEN = bool(getattr(sys, "frozen", False))
RESOURCE_DIR = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
BASE_DIR = Path(sys.executable).resolve().parent if IS_FROZEN else Path(__file__).resolve().parent
PROJECT_DIR = BASE_DIR.parent
WEB_DIR = RESOURCE_DIR / "web" if IS_FROZEN else PROJECT_DIR / "web"
ALGORITHM_WORKER = RESOURCE_DIR / "run_algorithm_worker.py" if IS_FROZEN else BASE_DIR / "run_algorithm_worker.py"
UPLOADS_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "output"
SESSION_FILE = BASE_DIR / "session.json"
REPORT_SCREENSHOT_ORDER = (
    "Left-Front-q",
    "Right-Front-q",
    "Left-Rear-q",
    "Right-Rear-q",
    "Left-Front-c",
    "Right-Front-c",
    "Left-Rear-c",
    "Right-Rear-c",
    "Left-Front-p",
    "Left-Front-p30",
    "Right-Front-p",
    "Right-Front-p30",
    "Left-Rear-p",
    "Left-Rear-p30",
    "Right-Rear-p",
    "Right-Rear-p30",
)

for directory in (UPLOADS_DIR, OUTPUT_DIR):
    directory.mkdir(parents=True, exist_ok=True)

logging.basicConfig(level=logging.INFO, format="%(asctime)s  %(levelname)-7s  %(message)s")
logger = logging.getLogger("wheelhouse_server")


def _detect_server_host() -> str:
    return os.environ.get("WHEELHOUSE_SERVER_HOST", "0.0.0.0")


def _detect_server_port() -> int:
    value = os.environ.get("WHEELHOUSE_SERVER_PORT", "8010")
    try:
        return int(value)
    except ValueError:
        return 8010


def _is_port_available(port: int) -> bool:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        return sock.connect_ex(("127.0.0.1", port)) != 0


def _select_server_port(preferred_port: int) -> int:
    if _is_port_available(preferred_port):
        return preferred_port
    if os.environ.get("WHEELHOUSE_SERVER_PORT"):
        raise SystemExit(f"端口 {preferred_port} 已被占用")
    for port in range(preferred_port + 1, preferred_port + 21):
        if _is_port_available(port):
            print(f"端口 {preferred_port} 已被占用，自动改用端口 {port}。访问地址: http://localhost:{port}")
            return port
    raise SystemExit(f"端口 {preferred_port}-{preferred_port + 20} 均被占用")


SERVER_HOST = _detect_server_host()
SERVER_PORT = _detect_server_port()
ACTIVE_SERVER_PORT = SERVER_PORT


class ConnectionManager:
    def __init__(self) -> None:
        self.active_connections: set[WebSocket] = set()

    async def connect(self, websocket: WebSocket) -> None:
        await websocket.accept()
        self.active_connections.add(websocket)

    def disconnect(self, websocket: WebSocket) -> None:
        self.active_connections.discard(websocket)

    async def broadcast(self, payload: dict[str, Any]) -> None:
        if not self.active_connections:
            return
        message = json.dumps(payload, ensure_ascii=False)
        stale: list[WebSocket] = []
        for websocket in self.active_connections:
            try:
                await websocket.send_text(message)
            except Exception:
                stale.append(websocket)
        for websocket in stale:
            self.disconnect(websocket)


manager = ConnectionManager()


class SessionState:
    def __init__(self) -> None:
        self.running = False
        self.completed = False
        self.last_result: dict[str, Any] | None = None
        self.last_error: str | None = None
        self.product_path: Path | None = None
        self.process_part_path: Path | None = None
        self.json_path: Path | None = None
        self.docx_report_path: Path | None = None
        self.uploaded_files: dict[str, str] = {}

    def to_dict(self) -> dict[str, Any]:
        return {
            "running": self.running,
            "completed": self.completed,
            "last_result": self.last_result,
            "last_error": self.last_error,
            "product_path": str(self.product_path) if self.product_path else None,
            "process_part_path": str(self.process_part_path) if self.process_part_path else None,
            "json_path": str(self.json_path) if self.json_path else None,
            "docx_report_path": str(self.docx_report_path) if self.docx_report_path else None,
            "uploaded_files": self.uploaded_files,
        }

    def save(self) -> None:
        SESSION_FILE.write_text(json.dumps(self.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8")

    def load(self) -> None:
        if not SESSION_FILE.exists():
            return
        try:
            data = json.loads(SESSION_FILE.read_text(encoding="utf-8"))
            self.running = False
            self.completed = bool(data.get("completed"))
            self.last_result = data.get("last_result")
            self.last_error = data.get("last_error")
            self.product_path = Path(data["product_path"]) if data.get("product_path") else None
            self.process_part_path = Path(data["process_part_path"]) if data.get("process_part_path") else None
            self.json_path = Path(data["json_path"]) if data.get("json_path") else None
            self.docx_report_path = Path(data["docx_report_path"]) if data.get("docx_report_path") else None
            self.uploaded_files = dict(data.get("uploaded_files") or {})
        except Exception:
            logger.warning("session.json 读取失败，使用空状态")


session = SessionState()
session.load()
catia_workflow_lock = asyncio.Lock()
current_task: asyncio.Task | None = None


async def emit_log(message: str, level: str = "info") -> None:
    logger.info(message)
    await manager.broadcast(
        {
            "type": "log",
            "data": {
                "time": datetime.now().strftime("%H:%M:%S"),
                "level": level,
                "message": message,
            },
        }
    )


async def broadcast_captured_log(message: str, level: str = "info") -> None:
    await manager.broadcast(
        {
            "type": "log",
            "data": {
                "time": datetime.now().strftime("%H:%M:%S"),
                "level": level,
                "message": message,
            },
        }
    )


async def emit_status() -> None:
    await manager.broadcast({"type": "status", "data": get_status_payload()})


def workflow_is_active() -> bool:
    return (current_task is not None and not current_task.done()) or catia_workflow_lock.locked()


def get_status_payload() -> dict[str, Any]:
    if session.running and not workflow_is_active():
        session.running = False
        session.save()
    return {
        "running": session.running,
        "completed": session.completed,
        "uploaded_files": session.uploaded_files,
        "saved_as_path": str(session.product_path) if session.product_path else None,
        "report_path": str(session.json_path) if session.json_path else None,
        "docx_report_path": str(session.docx_report_path) if session.docx_report_path else None,
        "last_error": session.last_error,
        "has_result_file": bool(session.product_path and session.product_path.exists()),
        "has_report_file": bool(session.json_path and session.json_path.exists()),
        "has_docx_report_file": bool(session.docx_report_path and session.docx_report_path.exists()),
    }


def sanitize_filename(filename: str, fallback: str) -> str:
    name = Path(filename).name
    name = re.sub(r"[^\w.\-\u4e00-\u9fff]+", "_", name, flags=re.UNICODE)
    return name or fallback


def normalize_upload_relative_path(raw_path: str, fallback: str) -> Path:
    value = (raw_path or "").replace("\\", "/").strip() or fallback
    path = Path(value)
    if path.is_absolute() or any(part in ("", ".", "..") for part in path.parts):
        raise HTTPException(status_code=400, detail=f"上传目录中包含非法相对路径: {raw_path}")
    return path


def save_upload_file(upload_file: UploadFile, target_path: Path) -> None:
    target_path.parent.mkdir(parents=True, exist_ok=True)
    with target_path.open("wb") as output:
        shutil.copyfileobj(upload_file.file, output)


def validate_upload_extension(filename: str, allowed: set[str]) -> None:
    suffix = Path(filename).suffix.casefold()
    if suffix not in allowed:
        raise HTTPException(status_code=400, detail=f"文件类型不正确: {filename}")


def parse_positive_float(value: str, label: str) -> float:
    try:
        number = float(value)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"{label} 必须是数字") from exc
    if number <= 0:
        raise HTTPException(status_code=400, detail=f"{label} 必须大于 0")
    return number


def collect_selected_measurement_keys(result: dict[str, Any]) -> list[str]:
    wheelhouse_inputs = result.get("wheelhouse_inputs") or []
    selected_keys: list[str] = []
    for item in wheelhouse_inputs:
        prefix = str((item or {}).get("measurement_prefix") or "").strip()
        if not prefix:
            continue
        for suffix in ("q", "c", "p", "p30"):
            selected_keys.append(f"{prefix}-{suffix}")
    if selected_keys:
        return selected_keys
    return [
        "Left-Front-q", "Right-Front-q", "Left-Rear-q", "Right-Rear-q",
        "Left-Front-c", "Right-Front-c", "Left-Rear-c", "Right-Rear-c",
        "Left-Front-p", "Left-Front-p30", "Right-Front-p", "Right-Front-p30",
        "Left-Rear-p", "Left-Rear-p30", "Right-Rear-p", "Right-Rear-p30",
    ]


def collect_regulation_distance_measurements(result: dict[str, Any]) -> dict[str, float]:
    measurements: dict[str, float] = {}
    selected_keys = set(collect_selected_measurement_keys(result))
    regulation_axis_part = result.get("regulation_axis_part") or {}
    for key, value in (regulation_axis_part.get("bbox_measurements") or {}).items():
        if key not in selected_keys:
            continue
        try:
            measurements[str(key)] = round(float(value), 2)
        except Exception:
            pass
    for item in regulation_axis_part.get("section_curve_results") or []:
        for key, value in (item.get("regulation_distance_measurements") or {}).items():
            if key not in selected_keys:
                continue
            try:
                measurements[str(key)] = round(float(value), 2)
            except Exception:
                pass
    return measurements


def format_mm(value: float | None) -> str:
    if value is None:
        return "N/A"
    return f"{float(value):.2f}"


def report_cell(text: str, bold: bool = False, fill: str | None = None, align: str = "center") -> str:
    text_xml = xml_escape(text)
    fill_xml = f"<w:shd w:fill=\"{fill}\"/>" if fill else ""
    bold_xml = "<w:b/>" if bold else ""
    return (
        "<w:tc><w:tcPr><w:tcW w:w=\"1800\" w:type=\"dxa\"/>"
        "<w:tcMar><w:top w:w=\"100\" w:type=\"dxa\"/><w:bottom w:w=\"100\" w:type=\"dxa\"/>"
        "<w:start w:w=\"120\" w:type=\"dxa\"/><w:end w:w=\"120\" w:type=\"dxa\"/></w:tcMar>"
        f"{fill_xml}<w:vAlign w:val=\"center\"/></w:tcPr>"
        f"<w:p><w:pPr><w:jc w:val=\"{align}\"/></w:pPr>"
        "<w:r><w:rPr><w:rFonts w:ascii=\"Microsoft YaHei\" w:hAnsi=\"Microsoft YaHei\" "
        "w:eastAsia=\"Microsoft YaHei\"/><w:sz w:val=\"20\"/>"
        f"{bold_xml}</w:rPr><w:t>{text_xml}</w:t></w:r></w:p></w:tc>"
    )


def report_row(values: list[str], bold: bool = False, fill: str | None = None, left_first: bool = False) -> str:
    cells = []
    for index, value in enumerate(values):
        cells.append(report_cell(value, bold=bold, fill=fill, align="left" if left_first and index == 0 else "center"))
    return "<w:tr>" + "".join(cells) + "</w:tr>"


def report_paragraph(text: str, style: str | None = None) -> str:
    ppr = ""
    if style:
        ppr = f"<w:pPr><w:pStyle w:val=\"{style}\"/></w:pPr>"
    return (
        f"<w:p>{ppr}<w:r><w:rPr><w:rFonts w:ascii=\"Microsoft YaHei\" "
        "w:hAnsi=\"Microsoft YaHei\" w:eastAsia=\"Microsoft YaHei\"/>"
        f"</w:rPr><w:t>{xml_escape(text)}</w:t></w:r></w:p>"
    )


def build_compliance_rows(
    measurements: dict[str, float],
    tire_radius: float,
    tire_width_y: float,
    measurement_keys: list[str] | None = None,
) -> list[dict[str, Any]]:
    checks = []
    keys = measurement_keys or [
        "Left-Front-q", "Right-Front-q", "Left-Rear-q", "Right-Rear-q",
        "Left-Front-c", "Right-Front-c", "Left-Rear-c", "Right-Rear-c",
        "Left-Front-p", "Left-Front-p30", "Right-Front-p", "Right-Front-p30",
        "Left-Rear-p", "Left-Rear-p30", "Right-Rear-p", "Right-Rear-p30",
    ]
    for key in keys:
        if key.endswith("-q"):
            threshold = tire_width_y
            op = ">"
            requirement = f"> {format_mm(threshold)} mm"
        elif key.endswith("-c"):
            threshold = 2.0 * tire_radius
            op = "<"
            requirement = f"< {format_mm(threshold)} mm"
        else:
            threshold = 30.0
            op = ">"
            requirement = f"> {format_mm(threshold)} mm"
        value = measurements.get(key)
        passed = value is not None and ((value > threshold) if op == ">" else (value < threshold))
        checks.append((key, key, threshold, value, op, requirement, passed))
    rows: list[dict[str, Any]] = []
    for key, item, threshold, value, op, requirement, passed in checks:
        rows.append(
            {
                "key": key,
                "item": item,
                "value": value,
                "requirement": requirement,
                "result": "PASS" if passed else "FAIL",
            }
        )
    return rows


def write_minimal_docx_report(
    output_path: Path,
    measurements: dict[str, float],
    compliance_rows: list[dict[str, Any]],
    tire_radius: float,
    tire_width_y: float,
    measurement_keys: list[str] | None = None,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    all_pass = all(row.get("result") == "PASS" for row in compliance_rows)
    overall_result = "PASS" if all_pass else "FAIL"
    info_table = (
        "<w:tbl><w:tblPr><w:tblW w:w=\"9000\" w:type=\"dxa\"/>"
        "<w:tblBorders><w:top w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:left w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:bottom w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:right w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:insideH w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:insideV w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/></w:tblBorders></w:tblPr>"
        + report_row(["报告类型", "车轮罩法规校核", "总体判定", overall_result], bold=False, fill="F2F4F7")
        + report_row(["轮胎半径 r", f"{format_mm(tire_radius)} mm", "轮胎Y向宽度 b", f"{format_mm(tire_width_y)} mm"], bold=False)
        + "</w:tbl>"
    )
    measurement_rows = "".join(
        report_row(
            [key, f"{format_mm(measurements.get(key))} mm" if measurements.get(key) is not None else "-"],
            left_first=True,
        )
        for key in (measurement_keys or [
            "Left-Front-q", "Right-Front-q", "Left-Rear-q", "Right-Rear-q",
            "Left-Front-c", "Right-Front-c", "Left-Rear-c", "Right-Rear-c",
            "Left-Front-p", "Left-Front-p30", "Right-Front-p", "Right-Front-p30",
            "Left-Rear-p", "Left-Rear-p30", "Right-Rear-p", "Right-Rear-p30",
        ])
        if key in measurements
    )
    table_xml = (
        "<w:tbl><w:tblPr><w:tblW w:w=\"9000\" w:type=\"dxa\"/>"
        "<w:tblBorders><w:top w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:left w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:bottom w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:right w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:insideH w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:insideV w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/></w:tblBorders></w:tblPr>"
        + report_row(["Measurement", "Value"], bold=True, fill="E8EEF5")
        + measurement_rows
        + "</w:tbl>"
    )
    comparison_table = (
        "<w:tbl><w:tblPr><w:tblW w:w=\"9000\" w:type=\"dxa\"/>"
        "<w:tblBorders><w:top w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:left w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:bottom w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:right w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:insideH w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/>"
        "<w:insideV w:val=\"single\" w:sz=\"4\" w:color=\"D0D5DD\"/></w:tblBorders></w:tblPr>"
        + report_row(["Item", "Value", "Requirement", "Result"], bold=True, fill="E8EEF5")
        + "".join(
            report_row([
                row["item"],
                f"{format_mm(row['value'])} mm",
                row["requirement"],
                row["result"],
            ], left_first=True)
            for row in compliance_rows
        )
        + "</w:tbl>"
    )
    document_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
        "<w:body>"
        + report_paragraph("车轮罩法规校核报告", "Title")
        + report_paragraph("Wheelhouse Regulation Verification Report", "Subtitle")
        + report_paragraph("报告信息", "Heading1")
        + info_table
        + report_paragraph(f"结论摘要：本次测量结果按 p > 30 mm、c < 2r、q > b 进行校核，总体判定为 {overall_result}。")
        + report_paragraph("车轮罩法规距离测量", "Heading1")
        + table_xml
        + report_paragraph("法规对比结果", "Heading1")
        + comparison_table
        + report_paragraph("判定规则", "Heading1")
        + report_paragraph("1. 两处 p 值需满足 > 30 mm。")
        + report_paragraph("2. c 值需满足 < 2r。")
        + report_paragraph("3. q 值需满足 > b。")
        + "<w:sectPr><w:pgSz w:w=\"11906\" w:h=\"16838\"/><w:pgMar w:top=\"1224\" w:right=\"1152\" w:bottom=\"1080\" w:left=\"1152\"/></w:sectPr>"
        "</w:body></w:document>"
    )
    content_types = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\"?>"
        "<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">"
        "<Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>"
        "<Default Extension=\"xml\" ContentType=\"application/xml\"/>"
        "<Override PartName=\"/word/document.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/>"
        "<Override PartName=\"/word/styles.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml\"/>"
        "</Types>"
    )
    rels = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\"?>"
        "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
        "<Relationship Id=\"rId1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" Target=\"word/document.xml\"/>"
        "</Relationships>"
    )
    document_rels = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\"?>"
        "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
        "<Relationship Id=\"rId1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles\" Target=\"styles.xml\"/>"
        "</Relationships>"
    )
    styles_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<w:styles xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
        "<w:style w:type=\"paragraph\" w:default=\"1\" w:styleId=\"Normal\"><w:name w:val=\"Normal\"/>"
        "<w:pPr><w:spacing w:after=\"120\" w:line=\"264\" w:lineRule=\"auto\"/></w:pPr>"
        "<w:rPr><w:rFonts w:ascii=\"Microsoft YaHei\" w:hAnsi=\"Microsoft YaHei\" w:eastAsia=\"Microsoft YaHei\"/>"
        "<w:sz w:val=\"21\"/><w:color w:val=\"1F2937\"/></w:rPr></w:style>"
        "<w:style w:type=\"paragraph\" w:styleId=\"Title\"><w:name w:val=\"Title\"/><w:basedOn w:val=\"Normal\"/>"
        "<w:pPr><w:jc w:val=\"center\"/><w:spacing w:after=\"80\"/></w:pPr>"
        "<w:rPr><w:rFonts w:ascii=\"Microsoft YaHei\" w:hAnsi=\"Microsoft YaHei\" w:eastAsia=\"Microsoft YaHei\"/>"
        "<w:b/><w:sz w:val=\"44\"/><w:color w:val=\"1F4D78\"/></w:rPr></w:style>"
        "<w:style w:type=\"paragraph\" w:styleId=\"Subtitle\"><w:name w:val=\"Subtitle\"/><w:basedOn w:val=\"Normal\"/>"
        "<w:pPr><w:jc w:val=\"center\"/><w:spacing w:after=\"240\"/></w:pPr>"
        "<w:rPr><w:rFonts w:ascii=\"Microsoft YaHei\" w:hAnsi=\"Microsoft YaHei\" w:eastAsia=\"Microsoft YaHei\"/>"
        "<w:sz w:val=\"20\"/><w:color w:val=\"667085\"/></w:rPr></w:style>"
        "<w:style w:type=\"paragraph\" w:styleId=\"Heading1\"><w:name w:val=\"heading 1\"/><w:basedOn w:val=\"Normal\"/>"
        "<w:pPr><w:spacing w:before=\"280\" w:after=\"120\"/><w:outlineLvl w:val=\"0\"/></w:pPr>"
        "<w:rPr><w:rFonts w:ascii=\"Microsoft YaHei\" w:hAnsi=\"Microsoft YaHei\" w:eastAsia=\"Microsoft YaHei\"/>"
        "<w:b/><w:sz w:val=\"30\"/><w:color w:val=\"2E74B5\"/></w:rPr></w:style>"
        "</w:styles>"
    )
    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", content_types)
        docx.writestr("_rels/.rels", rels)
        docx.writestr("word/_rels/document.xml.rels", document_rels)
        docx.writestr("word/styles.xml", styles_xml)
        docx.writestr("word/document.xml", document_xml)


def collect_report_screenshot_rows(result: dict[str, Any], result_dir: Path) -> list[dict[str, Any]]:
    screenshot_result = result.get("screenshot_result") or {}
    screenshots = screenshot_result.get("screenshots") or []
    by_key: dict[str, dict[str, Any]] = {}
    for row in screenshots:
        measurement_key = str(row.get("measurement_key") or "").strip()
        if not measurement_key:
            continue
        path_text = str(row.get("path") or row.get("screenshot_path") or row.get("output_path") or "").strip()
        if not path_text:
            continue
        path = Path(path_text)
        if not path.is_absolute():
            path = result_dir / path
        by_key[measurement_key] = {
            "measurement_key": measurement_key,
            "path": str(path),
            "status": row.get("status"),
            "category": row.get("category"),
            "annotation_name": row.get("annotation_name"),
        }
    rows: list[dict[str, Any]] = []
    for measurement_key in REPORT_SCREENSHOT_ORDER:
        row = by_key.get(measurement_key)
        if row and Path(row["path"]).is_file():
            rows.append(row)
    return rows


def append_screenshots_to_docx_report(report_path: Path, screenshot_rows: list[dict[str, Any]]) -> dict[str, Any]:
    if not screenshot_rows:
        return {
            "status": "skipped",
            "message": "未找到可写入 Word 报告的截图。",
            "count": 0,
            "success_count": 0,
            "screenshots": [],
        }
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
    except Exception as exc:
        return {
            "status": "failed",
            "message": f"无法加载 python-docx，未写入截图: {exc}",
            "count": len(screenshot_rows),
            "success_count": 0,
            "screenshots": screenshot_rows,
        }

    appended_rows: list[dict[str, Any]] = []
    failed_rows: list[dict[str, Any]] = []
    try:
        document = Document(str(report_path))
        available_width = None
        try:
            section = document.sections[-1]
            available_width = section.page_width - section.left_margin - section.right_margin
        except Exception:
            available_width = None

        for row in screenshot_rows:
            measurement_key = str(row.get("measurement_key") or "").strip()
            image_path = str(row.get("path") or "").strip()
            try:
                title = document.add_paragraph()
                title.paragraph_format.space_before = Pt(6)
                title.paragraph_format.space_after = Pt(2)
                title_run = title.add_run(measurement_key)
                title_run.bold = True
                title_run.font.size = Pt(12)

                image_paragraph = document.add_paragraph()
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_paragraph.paragraph_format.space_before = Pt(0)
                image_paragraph.paragraph_format.space_after = Pt(6)
                image_run = image_paragraph.add_run()
                compact_width = Inches(5.6)
                if available_width is not None:
                    image_run.add_picture(image_path, width=min(available_width, compact_width))
                else:
                    image_run.add_picture(image_path, width=compact_width)
                appended_rows.append(
                    {
                        "measurement_key": measurement_key,
                        "path": image_path,
                        "status": "success",
                    }
                )
            except Exception as exc:
                failed_rows.append(
                    {
                        "measurement_key": measurement_key,
                        "path": image_path,
                        "status": "failed",
                        "message": str(exc),
                    }
                )
        document.save(str(report_path))
    except Exception as exc:
        return {
            "status": "failed",
            "message": str(exc),
            "count": len(screenshot_rows),
            "success_count": len(appended_rows),
            "screenshots": appended_rows + failed_rows,
        }

    status = "success"
    if failed_rows:
        status = "partial_failed" if appended_rows else "failed"
    return {
        "status": status,
        "count": len(screenshot_rows),
        "success_count": len(appended_rows),
        "screenshots": appended_rows + failed_rows,
    }


def create_regulation_docx_report(result: dict[str, Any], tire_radius: float, tire_width_y: float) -> tuple[Path, dict[str, Any]]:
    run_output_paths = result.get("run_output_paths") or {}
    result_dir = Path(run_output_paths.get("result_dir") or OUTPUT_DIR)
    timestamp = run_output_paths.get("timestamp") or datetime.now().strftime("%Y%m%d_%H%M%S")
    measurements = collect_regulation_distance_measurements(result)
    measurement_keys = collect_selected_measurement_keys(result)
    compliance_rows = build_compliance_rows(measurements, tire_radius, tire_width_y, measurement_keys)
    report_path = result_dir / f"wheelhouse_regulation_report_{timestamp}.docx"
    write_minimal_docx_report(
        report_path,
        measurements,
        compliance_rows,
        tire_radius,
        tire_width_y,
        measurement_keys,
    )
    screenshot_rows = collect_report_screenshot_rows(result, result_dir)
    screenshot_append_result = append_screenshots_to_docx_report(report_path, screenshot_rows)
    report_data = {
        "tire_radius_r": tire_radius,
        "tire_width_y_b": tire_width_y,
        "measurements": measurements,
        "compliance_results": compliance_rows,
        "screenshot_append_result": screenshot_append_result,
        "docx_report_path": str(report_path),
    }
    result["regulation_report"] = report_data
    return report_path, report_data


def get_catia_application(start_if_missing: bool = False) -> Any:
    try:
        import win32com.client  # type: ignore
    except Exception as exc:
        raise RuntimeError("当前环境未安装 pywin32，无法检测 CATIA") from exc
    try:
        return win32com.client.GetActiveObject("CATIA.Application")
    except Exception:
        if not start_if_missing:
            raise RuntimeError("CATIA 未打开")
    try:
        catia = win32com.client.Dispatch("CATIA.Application")
        catia.Visible = True
        return catia
    except Exception as exc:
        raise RuntimeError(f"启动 CATIA 失败: {exc}") from exc


def first_existing_path(*paths: Any) -> Path | None:
    for item in paths:
        if item:
            path = Path(str(item))
            if path.exists():
                return path
    return None


def run_algorithm_sync(
    left_front_path: Path | None,
    right_front_path: Path | None,
    left_rear_path: Path | None,
    right_rear_path: Path | None,
    wheel_path: Path,
    loop: asyncio.AbstractEventLoop | None = None,
) -> dict[str, Any]:
    if not IS_FROZEN and not ALGORITHM_WORKER.exists():
        raise FileNotFoundError(f"未找到检测子进程文件: {ALGORITHM_WORKER}")

    result_file = Path(tempfile.gettempdir()) / f"wheelhouse_result_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}.json"
    env = os.environ.copy()
    env["PYTHONIOENCODING"] = "utf-8"
    def _arg(value: Path | None) -> str:
        return str(value) if value is not None else ""
    wheelhouse_args = [_arg(left_front_path), _arg(right_front_path), _arg(left_rear_path), _arg(right_rear_path)]
    command = (
        [sys.executable, "--wheelhouse-worker", *wheelhouse_args, str(wheel_path), str(result_file), str(OUTPUT_DIR)]
        if IS_FROZEN
        else [sys.executable, str(ALGORITHM_WORKER), *wheelhouse_args, str(wheel_path), str(result_file), str(OUTPUT_DIR)]
    )

    process = subprocess.Popen(
        command,
        cwd=str(BASE_DIR),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        encoding="utf-8",
        errors="replace",
        env=env,
    )

    output_tail: list[str] = []
    assert process.stdout is not None
    for raw_line in process.stdout:
        line = raw_line.rstrip("\r\n")
        if not line.strip():
            continue
        logger.info("[worker] %s", line)
        output_tail.append(line)
        output_tail = output_tail[-80:]
        if loop is not None:
            asyncio.run_coroutine_threadsafe(broadcast_captured_log(line), loop)

    return_code = process.wait()
    if result_file.exists():
        try:
            result = json.loads(result_file.read_text(encoding="utf-8"))
        finally:
            try:
                result_file.unlink()
            except Exception:
                pass
        if return_code != 0 and result.get("success", False):
            result["success"] = False
            result["error"] = f"检测子进程异常退出，退出码: {return_code}"
        return result

    detail = "\n".join(output_tail[-20:]).strip()
    error = f"检测子进程崩溃或未返回结果，退出码: {return_code}"
    if detail:
        error = f"{error}\n最后输出:\n{detail}"
    return {"success": False, "error": error, "worker_exit_code": return_code}


@asynccontextmanager
async def lifespan(_app: FastAPI):
    logger.info("Wheelhouse regulation server started")
    logger.info("访问地址: http://localhost:%s", ACTIVE_SERVER_PORT)
    yield


app = FastAPI(
    title="Wheelhouse Regulation Verification",
    description="CATIA 车轮罩法规校核前后端服务",
    version="1.0.0",
    lifespan=lifespan,
)


@app.get("/", response_class=HTMLResponse)
async def index():
    index_file = WEB_DIR / "index.html"
    if not index_file.exists():
        raise HTTPException(status_code=404, detail="web/index.html 不存在")
    return HTMLResponse(index_file.read_text(encoding="utf-8"), headers={"Cache-Control": "no-store"})


@app.get("/app.js")
async def app_js():
    return FileResponse(WEB_DIR / "app.js", media_type="application/javascript", headers={"Cache-Control": "no-store"})


@app.get("/style.css")
async def style_css():
    return FileResponse(WEB_DIR / "style.css", media_type="text/css", headers={"Cache-Control": "no-store"})


@app.post("/verify")
async def verify_wheelhouse(
    left_front_wheelhouse_file: UploadFile | None = File(None),
    right_front_wheelhouse_file: UploadFile | None = File(None),
    left_rear_wheelhouse_file: UploadFile | None = File(None),
    right_rear_wheelhouse_file: UploadFile | None = File(None),
    wheel_assembly_files: list[UploadFile] = File(...),
    wheel_assembly_root_path: str = Form(...),
    tire_radius_r: str = Form(...),
    tire_width_y_b: str = Form(...),
):
    global current_task
    if current_task is not None and current_task.done():
        current_task = None
    if current_task is not None and not current_task.done():
        raise HTTPException(status_code=409, detail="检测流程正在运行，请稍后再试")
    if catia_workflow_lock.locked():
        raise HTTPException(status_code=503, detail="检测任务进行中，请稍后再试")

    wheel_root_rel = normalize_upload_relative_path(wheel_assembly_root_path, "wheel_assembly.CATProduct")
    wheel_name = sanitize_filename(wheel_root_rel.name, "wheel_assembly.CATProduct")
    validate_upload_extension(wheel_name, {".catproduct"})
    if not wheel_assembly_files:
        raise HTTPException(status_code=400, detail="请上传车轮装配所在文件夹")
    tire_radius = parse_positive_float(tire_radius_r, "轮胎半径值 r")
    tire_width_y = parse_positive_float(tire_width_y_b, "轮胎Y向宽度值 b")

    wheelhouse_uploads = [
        ("left_front", left_front_wheelhouse_file, "Left_Front_Wheelhouse.CATPart"),
        ("right_front", right_front_wheelhouse_file, "Right_Front_Wheelhouse.CATPart"),
        ("left_rear", left_rear_wheelhouse_file, "Left_Rear_Wheelhouse.CATPart"),
        ("right_rear", right_rear_wheelhouse_file, "Right_Rear_Wheelhouse.CATPart"),
    ]
    provided_wheelhouse_uploads = [(slot, upload) for slot, upload, _default_name in wheelhouse_uploads if upload is not None]
    if not provided_wheelhouse_uploads:
        raise HTTPException(status_code=400, detail="请至少上传一个轮罩零件")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    task_dir = UPLOADS_DIR / timestamp
    task_dir.mkdir(parents=True, exist_ok=True)
    wheel_folder_path = task_dir / "wheel_assembly_folder"
    wheelhouse_paths: dict[str, Path] = {}
    wheelhouse_names: dict[str, str] = {}
    for slot, upload_file, default_name in wheelhouse_uploads:
        if upload_file is None:
            continue
        file_name = sanitize_filename(upload_file.filename or "", default_name)
        validate_upload_extension(file_name, {".catpart"})
        target_path = task_dir / f"{slot}_{file_name}"
        save_upload_file(upload_file, target_path)
        wheelhouse_paths[slot] = target_path
        wheelhouse_names[slot] = file_name
    saved_wheel_paths: list[Path] = []
    for upload_file in wheel_assembly_files:
        rel_path = normalize_upload_relative_path(
            upload_file.filename or "",
            sanitize_filename(upload_file.filename or "", "wheel_file"),
        )
        target_path = wheel_folder_path / rel_path
        save_upload_file(upload_file, target_path)
        saved_wheel_paths.append(target_path)
    wheel_path = wheel_folder_path / wheel_root_rel
    if not wheel_path.exists():
        matching = [path for path in saved_wheel_paths if path.name.casefold() == wheel_root_rel.name.casefold()]
        if len(matching) == 1:
            wheel_path = matching[0]
        else:
            raise HTTPException(status_code=400, detail=f"上传文件夹中未找到所选根装配: {wheel_assembly_root_path}")

    async with catia_workflow_lock:
        current_task = asyncio.current_task()
        session.running = True
        session.completed = False
        session.last_error = None
        session.last_result = None
        session.product_path = None
        session.process_part_path = None
        session.json_path = None
        session.docx_report_path = None
        session.uploaded_files = {
            **{slot: name for slot, name in wheelhouse_names.items()},
            "wheel_assembly_root": str(wheel_path),
            "wheel_assembly_root_relative_path": str(wheel_root_rel),
            "wheel_assembly_file_count": len(saved_wheel_paths),
        }
        session.save()
        await emit_status()
        await emit_log("开始车轮罩法规校核")
        for slot, name in wheelhouse_names.items():
            await emit_log(f"{slot}: {name}")
        await emit_log(f"车轮装配根节点: {wheel_root_rel}")
        await emit_log(f"车轮装配文件夹文件数: {len(saved_wheel_paths)}")

        try:
            loop = asyncio.get_running_loop()
            result = await loop.run_in_executor(
                None,
                run_algorithm_sync,
                wheelhouse_paths.get("left_front"),
                wheelhouse_paths.get("right_front"),
                wheelhouse_paths.get("left_rear"),
                wheelhouse_paths.get("right_rear"),
                wheel_path,
                loop,
            )
            session.last_result = result
            session.completed = bool(result.get("success"))
            regulation_axis_part = result.get("regulation_axis_part") or {}
            run_output_paths = result.get("run_output_paths") or {}
            session.product_path = first_existing_path(result.get("product_path"))
            session.process_part_path = first_existing_path(regulation_axis_part.get("part_save_path"))
            session.json_path = first_existing_path(run_output_paths.get("default_json_result_path"))

            if not result.get("success"):
                session.last_error = str(result.get("error", "未知错误"))
                await emit_log(f"校核失败: {session.last_error}", level="error")
                raise HTTPException(status_code=500, detail=session.last_error)
            if not session.product_path:
                raise HTTPException(status_code=500, detail="校核完成但未生成总成 CATProduct")
            await emit_log("[11] 保存校核结果")
            session.docx_report_path, report_data = create_regulation_docx_report(
                result,
                tire_radius,
                tire_width_y,
            )
            session.last_result = result
            await emit_log("[11] 保存校核结果完成", level="success")

            await emit_log("车轮罩法规校核完成", level="success")
            await emit_log(f"总成文件: {session.product_path}")
            if session.process_part_path:
                await emit_log(f"过程 Part: {session.process_part_path}")
            if session.json_path:
                await emit_log(f"结果 JSON: {session.json_path}")
            if session.docx_report_path:
                await emit_log(f"法规校核报告: {session.docx_report_path}")
            return {
                "ok": True,
                "result_filename": session.product_path.name,
                "process_filename": session.process_part_path.name if session.process_part_path else None,
                "report_filename": session.json_path.name if session.json_path else None,
                "docx_report_filename": session.docx_report_path.name if session.docx_report_path else None,
                "result_download_url": "/api/download-result",
                "process_download_url": "/api/download-process",
                "report_download_url": "/api/download-report",
                "docx_report_download_url": "/api/download-docx-report",
                "regulation_report": report_data,
            }
        except HTTPException:
            raise
        except Exception as exc:
            session.last_error = str(exc)
            session.last_result = {"success": False, "error": str(exc), "traceback": traceback.format_exc()}
            await emit_log(f"校核异常: {exc}", level="error")
            raise HTTPException(status_code=500, detail=f"检测失败: {exc}") from exc
        finally:
            session.running = False
            session.save()
            current_task = None
            await emit_status()


@app.post("/api/reset")
async def reset():
    global current_task
    if current_task is not None and current_task.done():
        current_task = None
    if current_task is not None and not current_task.done() or catia_workflow_lock.locked():
        raise HTTPException(status_code=409, detail="检测流程正在运行，暂不能重置")
    session.running = False
    session.completed = False
    session.last_error = None
    session.last_result = None
    session.product_path = None
    session.process_part_path = None
    session.json_path = None
    session.docx_report_path = None
    session.uploaded_files = {}
    session.save()
    await emit_log("已重置，请重新上传轮罩零件和车轮装配文件夹")
    await emit_status()
    return {"ok": True}


@app.get("/api/status")
async def status():
    return get_status_payload()


@app.get("/api/catia-status")
async def catia_status():
    try:
        catia = get_catia_application(start_if_missing=False)
        return {"running": True, "visible": bool(getattr(catia, "Visible", False)), "message": "CATIA 已打开"}
    except Exception as exc:
        return {"running": False, "visible": False, "message": str(exc)}


@app.post("/api/open-catia")
async def open_catia():
    try:
        catia = get_catia_application(start_if_missing=True)
        try:
            catia.Visible = True
        except Exception:
            pass
        await emit_log("CATIA 已打开或已连接", level="success")
        return {"ok": True, "running": True, "visible": bool(getattr(catia, "Visible", False)), "message": "CATIA 已打开或已连接"}
    except Exception as exc:
        await emit_log(f"CATIA 打开失败: {exc}", level="error")
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.get("/health")
async def health():
    return {"status": "ok"}


@app.get("/api/result")
async def result():
    if session.last_result is None:
        raise HTTPException(status_code=404, detail="暂无检测结果")
    return session.last_result


@app.get("/api/download-result")
async def download_result():
    if not session.product_path or not session.product_path.exists():
        raise HTTPException(status_code=404, detail="暂无可下载的总成文件")
    return FileResponse(session.product_path, media_type="application/octet-stream", filename=session.product_path.name)


@app.get("/api/download-process")
async def download_process():
    if not session.process_part_path or not session.process_part_path.exists():
        raise HTTPException(status_code=404, detail="暂无可下载的过程 Part")
    return FileResponse(session.process_part_path, media_type="application/octet-stream", filename=session.process_part_path.name)


@app.get("/api/download-report")
async def download_report():
    if not session.json_path or not session.json_path.exists():
        raise HTTPException(status_code=404, detail="暂无可下载的 JSON 结果文件")
    return FileResponse(session.json_path, media_type="application/json", filename=session.json_path.name)


@app.get("/api/download-docx-report")
async def download_docx_report():
    if not session.docx_report_path or not session.docx_report_path.exists():
        raise HTTPException(status_code=404, detail="暂无可下载的法规校核报告")
    return FileResponse(
        session.docx_report_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=session.docx_report_path.name,
    )


@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await manager.connect(websocket)
    await websocket.send_text(json.dumps({"type": "status", "data": get_status_payload()}, ensure_ascii=False))
    try:
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        manager.disconnect(websocket)
    except Exception:
        manager.disconnect(websocket)


if WEB_DIR.exists():
    lib_dir = WEB_DIR / "lib"
    if lib_dir.exists():
        app.mount("/lib", StaticFiles(directory=str(lib_dir)), name="lib")


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--wheelhouse-worker":
        from run_algorithm_worker import main as worker_main

        sys.argv = [sys.argv[0], *sys.argv[2:]]
        raise SystemExit(worker_main())

    import uvicorn

    selected_port = _select_server_port(SERVER_PORT)
    ACTIVE_SERVER_PORT = selected_port
    uvicorn.run(app, host=SERVER_HOST, port=selected_port, reload=False, log_level="info")
